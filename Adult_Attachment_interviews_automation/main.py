import re
import os
import glob
import docx2txt
import sklearn
import pandas as pd
import numpy as np
import docx
import nltk
import matplotlib.pyplot as plt
from simplemma import simple_tokenizer
from hvplot import pandas
from wordcloud import WordCloud
from sklearn.model_selection import train_test_split,StratifiedKFold, KFold
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.svm import LinearSVC
from sklearn.ensemble import RandomForestClassifier
from sklearn.neural_network import MLPClassifier
from sklearn.naive_bayes import GaussianNB
import liwc
from collections import Counter
from sklearn.compose import ColumnTransformer
from sklearn.svm import SVC
from sklearn.impute import SimpleImputer
from sklearn.cluster import KMeans
from sklearn.metrics import classification_report, confusion_matrix
from sklearn.tree import DecisionTreeClassifier
from sklearn.model_selection import cross_val_score
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.dummy import DummyClassifier
import warnings

warnings.filterwarnings('ignore')

# Reading and preparing the secure/dismissing/preoccupied Dataframe 'df1'
df1 = pd.read_excel(
     os.path.join(r"C:\Users\Orelson\PycharmProjects\testing\Daten_Psychosomatische_Medizin/", 'AAI_Kategorien_subsamples.xlsx'),
     engine='openpyxl',
)
df1 = df1.drop(df1.index[0])
df1 = df1.drop(columns=df1.columns[1])
df1 = df1.drop(columns=df1.columns[2])
df1.columns = ["secure", "dismissing", "preoccupied"]

##################### getting rid of questions part in the files ##############################
# for file in glob.glob(os.path.join(r"C:\Users\Orelson\PycharmProjects\testing\Daten_Psychosomatische_Medizin/",'*.doc')):
#    if file != r"C:\Users\Orelson\PycharmProjects\testing\Daten_Psychosomatische_Medizin\~$I -  8109.docx" :
#    #if file == r"C:\Users\Orelson\PycharmProjects\testing\Daten_Psychosomatische_Medizin\8109.doc":
#     document1 = docx.Document(file)
#     #print(len(document.paragraphs))
#     for paragraph in document1.paragraphs:
#         for run in paragraph.runs:
#             if run.bold== True:
#
#                run.text = ""
#                #print(run.text)
#     document1.save(file)

# Reading and preparing the Wordx files containing the Interviews 'df2'
text = []
ids = []
for file in glob.glob(os.path.join(r"C:\Users\Orelson\PycharmProjects\testing\Daten_Psychosomatische_Medizin/",'*.doc')):
   if file != r"C:\Users\Orelson\PycharmProjects\testing\Daten_Psychosomatische_Medizin\~$I -  8109.docx" :
    document = docx.Document(file)
    text2 = docx2txt.process(file)
    text.append(text2)
    id1 = os.path.basename(file)
    id1 = id1[:-4]
    ids.append(id1)


df2 = pd.DataFrame(
    {'text': text,
     'id': ids
    })
df2 = df2.replace(r'\n', '', regex=True)    #replacing '/n' with empty spaces

# Merging df1 into df2
j = 0
categorie = []

while j < len(df2):
     if df2["id"][j] in df1["secure"].to_string():
            categorie.append("secure")
     elif df2["id"][j] in df1["dismissing"].to_string():
            categorie.append("dismissing")
     elif df2["id"][j] in df1["preoccupied"].to_string():
            categorie.append("preoccupied")
     else :
         categorie.append("nan")
     j += 1
df2["categorie"] = categorie

print(df2["categorie"].value_counts())
# Text_Data Preprocessing/ cleaning
corpus = []

for i in range(0, len(df2)):
   clean_data = re.sub("[!@#$%^&*(){};:,.<>?`_+]", ' ', df2['text'][i])
   clean_data = clean_data.lower()
   # clean_data = simple_tokenizer(clean_data)
   # clean_data = ' '.join(clean_data)
   corpus.append(clean_data)
   df2['text'][i] = clean_data
#print(corpus)


# Word Frequency of most common words  // Descriptive analysis
from bokeh.models import NumeralTickFormatter
word_freq = pd.Series(" ".join(df2["text"]).split()).value_counts()
word_freq[1:40].rename("Word frequency of most common words in comments").hvplot.bar(
    rot=45).opts(width=700, height=400, yformatter=NumeralTickFormatter(format="0,0"))

#TODO: use of the pipeline and grid search functionality in sci-kit learn totune the parameter using pipeline and "GridSearchCV"

######################################################## LIWC  #######################################################

parse, category_names = liwc.load_token_parser('LIWC2015_GermanDictionary.dic')

def tokenize(text):
    for match in re.finditer(r'\w+', text, re.UNICODE):
        yield match.group(0)

def counting(cat_tokens):
    cat_counts = Counter(category for token in cat_tokens for category in parse(token))
    return cat_counts


################################# classification / Features model  ###################################

word_count = []
you_total = []
she_he = []
hear = []
feel = []
negation = []
assent = []
conj = []
anx = []
for i in range(0, len(df2)):
    all_tokens = tokenize(df2['text'][i])
    all_counts = counting(all_tokens)
    word_count.append(all_counts['function'])
    you_total.append(all_counts['you_total'])
    she_he.append(all_counts['shehe'])
    hear.append(all_counts['hear'])
    feel.append(all_counts['feel'])
    negation.append(all_counts['negate'])
    assent.append(all_counts['assent'])
    conj.append(all_counts['conj'])
    anx.append(all_counts['anx'])

df2["word_count"] = word_count
df2["you_total"] = you_total
df2["she_he"] = she_he
df2["hear"] = hear
df2["feel"] = feel
df2["negate"] = negation
df2["assent"] = assent
df2["conj"] = conj
df2["anx"] = anx
print(df2)

X = df2[['text', 'word_count', 'you_total', 'she_he', 'hear']]
X1 = df2[[ 'word_count', 'you_total', 'she_he', 'hear','feel','negate','assent','conj','anx']]
X2 = df2[[ 'word_count', 'feel', 'assent','conj']]
X3 = df2[['text']]
X4 = df2[['you_total','feel', 'assent','conj']]
X5 = df2[['you_total','feel','conj']]
X6 = df2[['text', 'you_total','feel', 'assent','conj']]
y = df2['categorie']

# Encoding the Dependent Variable
from sklearn.preprocessing import LabelEncoder
le = LabelEncoder()
y = le.fit_transform(y)
#print(y)


X_train, X_test, y_train, y_test = train_test_split(X4, y, stratify=y,random_state=1, test_size=0.25)

#X
# transformer = ColumnTransformer(transformers = [
#     ('vectorizer', TfidfVectorizer(), 'text')
# ], remainder='passthrough')

 #X1 and X4
t = [('feel', SimpleImputer(strategy='median'), [0])]
transformer = ColumnTransformer(transformers=t , remainder='passthrough')

#X2
# t = [('you_total', SimpleImputer(strategy='median'), [0])]
# transformer = ColumnTransformer(transformers=t , remainder='passthrough')

X_vec_train = transformer.fit_transform(X_train)
X_vec_test = transformer.transform(X_test)


# # Classification models to test
parameters={"n_neighbors": range(1, 50)}
classifiers = [
    DummyClassifier(random_state=1),
    GaussianNB(),
    LogisticRegression(solver="saga", random_state=1),
    LinearSVC(random_state=1),
    RandomForestClassifier(random_state=1),
    DecisionTreeClassifier(criterion = 'entropy', random_state = 1),
    KMeans(n_clusters=3, init = 'k-means++', random_state=1),
    MLPClassifier(
        random_state=1,
        solver="adam",
        hidden_layer_sizes=(12, 12, 12),
        activation="relu",
        early_stopping=True,
        n_iter_no_change=1,
    ),
]
# get names of the objects in list (too lazy for c&p...)
names = [re.match(r"[^\(]+", name.__str__())[0] for name in classifiers]
print(f"Classifiers to test: {names}")

# test all classifiers and save pred. results on test data
results = {}
for name, clf in zip(names, classifiers):
    print(f"Training classifier: {name}")
    clf.fit(X_vec_train, y_train)
    prediction = clf.predict(X_vec_test)
    #report = sklearn.metrics.classification_report(y_test, prediction, labels=np.unique(prediction))
    report = sklearn.metrics.classification_report(y_test, prediction)
    results[name] = report



#print(results)

# Prediction results
for k, v in results.items():
    print(f"Results for {k}:")
    print(f"{v}\n")

##########################################
classifier = RandomForestClassifier(random_state=1)


# Visualising the Training set results using ANOVA f-test
from sklearn.feature_selection import SelectKBest
from sklearn.feature_selection import f_classif
from sklearn.feature_selection import f_regression
from sklearn.feature_selection import mutual_info_classif

# feature selection
def select_features(X_ttrain, y_ttrain, X_ttest):
    #fs = SelectKBest(score_func=f_classif, k='all')
    #fs = SelectKBest(score_func=mutual_info_classif, k='all')
    fs = SelectKBest(score_func=f_regression, k='all')
    fs.fit(X_train, y_train)
    X_train_fs = fs.transform(X_train)
    X_test_fs = fs.transform(X_test)
    return X_train_fs, X_test_fs, fs

# feature selection
X_train_fs, X_test_fs, fs = select_features(X_vec_train, y_train, X_vec_test)
# what are scores for the features
for i in range(len(fs.scores_)):
    print('Feature %d: %f' % (i, fs.scores_[i]))
# plot the scores
plt.bar([i for i in range(len(fs.scores_))], fs.scores_)
plt.show()

#cross val
from sklearn.model_selection import cross_val_score, cross_val_predict
from sklearn import metrics
# Perform n-fold cross validation
model = RandomForestClassifier(random_state=1)
#model = DummyClassifier(random_state=1)
scores = cross_val_score(model, X4, y, cv=10,scoring='accuracy')
print( 'Cross-validated scores:', scores)
print("Accuracy of Model with Cross Validation is:",scores.mean() * 100 , '%')

## determinate the number of features to use to get the best precision
# from numpy import mean
# from numpy import std
# from pandas import read_csv
# from sklearn.model_selection import cross_val_score
# from sklearn.model_selection import RepeatedStratifiedKFold
# from sklearn.feature_selection import SelectKBest
# from sklearn.feature_selection import f_classif
# from sklearn.linear_model import LogisticRegression
# from sklearn.pipeline import Pipeline
#
# # evaluate a give model using cross-validation
# def evaluate_model(model, X, y):
#     cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=3, random_state=1)
#     scoress = cross_val_score(model, X, y, scoring='accuracy', cv=cv, n_jobs=-1, error_score='raise')
#     return scoress
#
# # define number of features to evaluate
# num_features = [i+1 for i in range(X.shape[1])]
# # enumerate each number of features
# results = list()
# for k in num_features:
# # create pipeline
#     model = RandomForestClassifier(random_state=1)
#     fs = SelectKBest(score_func=f_classif, k=k)
#     pipeline = Pipeline(steps=[('anova',fs), ('lr', model)])
# # evaluate the model
#     scoress = evaluate_model(pipeline, X1, y)
#     results.append(scoress)
# # summarize the results
#     print('>%d %.3f (%.3f)' % (k, mean(scoress), std(scoress)))
# # plot model performance for comparison
# plt.boxplot(results, labels=num_features, showmeans=True)
# plt.show()

###################################### Bert model

# from transformers import BertTokenizer, TFBertModel
#
# tokenizer = BertTokenizer.from_pretrained("bert-base-german-cased")
# PATH_GDRIVE_TMP = r"C:\Users\Orelson\Desktop\New folder"
# MAXLEN = 19200
#
# def preprocess_text(df2):
#     """ take texts and prepare as input features for BERT
#     """
#     input_ids = []
#     # For every sentence...
#     for comment in df2:
#         encoded_sent = tokenizer.encode_plus(
#             text=comment,
#             add_special_tokens=True,  # Add `[CLS]` and `[SEP]`
#             max_length=MAXLEN,  # Max length to truncate/pad
#             pad_to_max_length=True,  # Pad sentence to max length
#             return_attention_mask=False,  # attention mask not needed for our task
#         )
#         # Add the outputs to the lists
#         input_ids.append(encoded_sent.get("input_ids"))
#     return input_ids
#
# # comment = ["Ich liebe data-dive.com und meine Katze."]
# # input_ids = preprocess_text(comment)
# # print("Comment: ", comment)
# # print("Tokenized Comment: ", tokenizer.convert_ids_to_tokens(input_ids[0])[0:20])
# # print("Token IDs: ", input_ids[0][0:20])
#
# import pickle
#
# input_ids = preprocess_text(df2["text"])
# # tokenization takes quite long
# # we can save the result and load it quickly via pickle
# pickle.dump(input_ids, open(PATH_GDRIVE_TMP + "/input_ids.pkl", "wb"))
# # input_ids = pickle.load(open(PATH_GDRIVE_TMP+"/input_ids.pkl", "rb"))
#
# # Sample data for cross validation
# train_ids, test_ids, train_labels, test_labels = train_test_split(
#     input_ids, df2["categorie"], random_state=1, test_size=0.25, shuffle=True
# )
# print(f"Train set: {len(train_ids)}\nTest set: {len(test_ids)}")
#
# # Set Model Parameters
# MAXLEN = MAXLEN
# BATCH_SIZE_PER_REPLICA = 16
# BATCH_SIZE = BATCH_SIZE_PER_REPLICA * strategy.num_replicas_in_sync
# EPOCHS = 8
# LEARNING_RATE = 1e-5
# DATA_LENGTH = len(df2)